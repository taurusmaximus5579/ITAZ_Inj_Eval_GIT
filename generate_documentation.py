from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

filename = "docs/ITAZ_Inj_Eval_GIT_Dokumentation.docx"

doc = Document()

def add_heading(text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    if bold:
        run = p.add_run(text)
        run.bold = True
    else:
        p.add_run(text)
    return p


doc.add_heading('ITAZ Injection Evaluation Tool - Dokumentation', level=0)

doc.add_paragraph('Dies ist die technische Dokumentation des Python-Tools zur Auswertung von Injektordaten. Die Dokumentation beschreibt die Architektur, den Datenfluss, die Modulverantwortung und die Benutzung der GUI.')

add_heading('1. Ziel und Übersicht', level=1)
doc.add_paragraph('Ziel des Projekts ist es, die Auswertung von Injektordaten modular und wartbar als Python-Anwendung bereitzustellen. Die Anwendung kann CSV-Dateien einlesen, Signale vorverarbeiten, mehrere Auswertungen durchführen und Ergebnisse als Excel und interaktive Plots ausgeben.')

doc.add_paragraph('Der Gesamtablauf gliedert sich in folgende Bereiche:')
doc.add_paragraph('- Konfiguration und Validierung', style='List Bullet')
doc.add_paragraph('- Dateiimport und Rohdatenverarbeitung', style='List Bullet')
doc.add_paragraph('- Signalaufbereitung und Interpolation', style='List Bullet')
doc.add_paragraph('- Auswertungen (ICS, Nadelhub, Injektionsrate, Gain, Rate-Down, Shot2Shot)', style='List Bullet')
doc.add_paragraph('- Plot-Erzeugung und Ergebnis-Export', style='List Bullet')

add_heading('2. Projektstruktur', level=1)
doc.add_paragraph('Die Anwendung ist in mehrere Module unterteilt. Die wichtigsten Dateien und Ordner sind:')
doc.add_paragraph('- ITAZ_Inj_Eval_GUI_github.py: GUI-Einstiegspunkt mit Benutzerinteraktion.', style='List Bullet')
doc.add_paragraph('- ITAZ_Inj_Eval_208.py: Haupt-Orchestrator / Pipeline.', style='List Bullet')
doc.add_paragraph('- config.py: Konfigurationsmodell und Validierung.', style='List Bullet')
doc.add_paragraph('- io_utils.py: Dateisuche und CSV-Einlesen.', style='List Bullet')
doc.add_paragraph('- preprocessing.py: Rohdaten in standardisierte Signalstrukturen transformieren.', style='List Bullet')
doc.add_paragraph('- analysis/: Ordner mit den fachlichen Auswertungen.', style='List Bullet')
doc.add_paragraph('- plotting.py: Diagramm-Erstellung und Bildausgabe.', style='List Bullet')
doc.add_paragraph('- export.py: Excel-Export von Ergebnistabellen.', style='List Bullet')
doc.add_paragraph('- image_utils.py: In-Memory-Bildpuffer, um Festplattenspeicher zu sparen.', style='List Bullet')
doc.add_paragraph('- Fkt/: Alte Spezialfunktionen und Plot-Hilfen.', style='List Bullet')

add_heading('3. Module und Verantwortlichkeiten', level=1)
add_heading('3.1 GUI', level=2)
doc.add_paragraph('Datei: ITAZ_Inj_Eval_GUI_github.py')
doc.add_paragraph('Aufgaben:')
doc.add_paragraph('- Auswahl von CSV-Dateien', style='List Bullet')
doc.add_paragraph('- Eingabe von Parametern für Gas, Prüfkammer, Interpolation und Auswertung', style='List Bullet')
doc.add_paragraph('- Steuerung der Analyse', style='List Bullet')
doc.add_paragraph('- Anzeige von Logs, Bild-Tabs und Shot-Log-Tabellen', style='List Bullet')
doc.add_paragraph('- Steuerung des Bildspeicherverhaltens über Checkboxen', style='List Bullet')

add_heading('3.2 Pipeline', level=2)
doc.add_paragraph('Datei: ITAZ_Inj_Eval_208.py')
doc.add_paragraph('Aufgaben:')
doc.add_paragraph('- Validierung der Konfiguration', style='List Bullet')
doc.add_paragraph('- Erkennen und Einlesen der Messdateien', style='List Bullet')
doc.add_paragraph('- Weitergabe von Signalen und Ergebnissen an die Auswertungsfunktionen', style='List Bullet')
doc.add_paragraph('- Erstellung der Ergebnis- und Bildordner', style='List Bullet')

add_heading('3.3 Konfiguration', level=2)
doc.add_paragraph('Datei: config.py')
doc.add_paragraph('Aufgaben:')
doc.add_paragraph('- Definition der Standardwerte', style='List Bullet')
doc.add_paragraph('- Validierung der GUI-Eingaben', style='List Bullet')
doc.add_paragraph('- Bereitstellung der Auswertungsparameter an die Pipeline', style='List Bullet')

add_heading('3.4 Datei- und Datenimport', level=2)
doc.add_paragraph('Datei: io_utils.py')
doc.add_paragraph('Aufgaben:')
doc.add_paragraph('- Erkennen von CSV-Dateien', style='List Bullet')
doc.add_paragraph('- Laden der Messdaten in pandas-Strukturen', style='List Bullet')
doc.add_paragraph('- Lesen und Normieren von shot_log.csv', style='List Bullet')

add_heading('3.5 Vorverarbeitung', level=2)
doc.add_paragraph('Datei: preprocessing.py')
doc.add_paragraph('Aufgaben:')
doc.add_paragraph('- Erzeugen einer einheitlichen Signalstruktur', style='List Bullet')
doc.add_paragraph('- Interpolation auf eine gemeinsame Zeitbasis', style='List Bullet')
doc.add_paragraph('- Normierung und Umrechnung der Sensordaten', style='List Bullet')

add_heading('3.6 Analyse-Module', level=2)
doc.add_paragraph('Ordner: analysis/')
doc.add_paragraph('Aufgaben:')
doc.add_paragraph('- ICS-Auswertung: Plateau-Erkennung, Boost/Hold/Zero-Bereiche, ICS_ON-Zeiten', style='List Bullet')
doc.add_paragraph('- Needle-Lift-Analyse: Hubzeiten, Integrale und Histogramme', style='List Bullet')
doc.add_paragraph('- Injektionsrate: Rate-, Masse- und Integralauswertung', style='List Bullet')
doc.add_paragraph('- Gain-Kurve: Massendaten vs. ICS_ON mit Regression', style='List Bullet')
doc.add_paragraph('- Rate-Down-Auswertung: Druck vs. Masse-Regression', style='List Bullet')
doc.add_paragraph('- Shot2Shot: Shot-Statistiken, Scatter-Matrix und Trendauswertung', style='List Bullet')

add_heading('3.7 Plotting', level=2)
doc.add_paragraph('Datei: plotting.py und Fkt/FigDict02.py')
doc.add_paragraph('Aufgaben:')
doc.add_paragraph('- Erzeugen von interaktiven HTML-Plots', style='List Bullet')
doc.add_paragraph('- Erzeugen von Matplotlib-Plots für die GUI-Anzeige', style='List Bullet')
doc.add_paragraph('- Auswahl zwischen Festplatten-Speicherung oder In-Memory-Cache', style='List Bullet')

doc.add_heading('3.8 Export', level=2)
doc.add_paragraph('Datei: export.py')
doc.add_paragraph('Aufgaben:')
doc.add_paragraph('- Excel-Dateien mit Ergebnis-Tabellen erzeugen', style='List Bullet')
doc.add_paragraph('- Sammeln aller Auswertungsergebnisse in strukturierter Form', style='List Bullet')

add_heading('4. Bedienung', level=1)
doc.add_paragraph('Schritte zur Benutzung der Anwendung:')
doc.add_paragraph('- Starten Sie die Anwendung über ITAZ_Inj_Eval_GUI_github.py.', style='List Bullet')
doc.add_paragraph('- Wählen Sie die gewünschten CSV-Dateien aus.', style='List Bullet')
doc.add_paragraph('- Wählen Sie optional Rohdatenplot, Excel-Export, Bildspeicherung und Auswertungen aus.', style='List Bullet')
doc.add_paragraph('- Starten Sie die Analyse mit "Analyse starten".', style='List Bullet')
doc.add_paragraph('- Nachdem die Analyse abgeschlossen ist, werden die Plots und Ergebnisse angezeigt.', style='List Bullet')

add_heading('5. Ergebnisse und Ausgabe', level=1)
doc.add_paragraph('Die Anwendung erzeugt im Ordner der Messdaten typischerweise folgende Ausgaben:')
doc.add_paragraph('- Results/: Ergebnis-HTML und strukturierte Tabellen', style='List Bullet')
doc.add_paragraph('- Bilder/: Auswertungsgrafiken und Diagramme', style='List Bullet')
doc.add_paragraph('- Excel-Dateien: Exportierte Ergebnisse und Statistikdaten', style='List Bullet')

doc.add_heading('6. Besonderheiten', level=1)
doc.add_paragraph('Aktuelle Eigenschaften:')
doc.add_paragraph('- Interaktive HTML-Plotlegende: Ein Klick hebt eine Kurve hervor und dimmt die anderen.', style='List Bullet')
doc.add_paragraph('- Der Cursor zeigt bei aktivierter Auswahl nur die Werte der gewählten Kurve an.', style='List Bullet')
doc.add_paragraph('- Bilder können wahlweise nur im Speicher gehalten werden, um Festplattenplatz zu sparen.', style='List Bullet')

add_heading('7. Erweiterungsmöglichkeiten', level=1)
doc.add_paragraph('Mögliche Verbesserungen:')
doc.add_paragraph('- Zusätzliche Auswertungen für Gaswechsel oder Temperaturverläufe', style='List Bullet')
doc.add_paragraph('- Erweiterung der GUI um Filter und Dateigruppierung', style='List Bullet')
doc.add_paragraph('- Verbesserte Reportgenerierung mit Zusammenfassungen und Parametern', style='List Bullet')
doc.add_paragraph('- Logik für dynamische Speicherbereinigung und Temp-Dateien', style='List Bullet')

add_heading('8. Technische Details', level=1)
doc.add_paragraph('Wichtige Bibliotheken:')
doc.add_paragraph('- pandas: Datenstrukturen und CSV-Einlese', style='List Bullet')
doc.add_paragraph('- numpy: numerische Berechnungen', style='List Bullet')
doc.add_paragraph('- matplotlib: Bildgenerierung für GUI und Auswertung', style='List Bullet')
doc.add_paragraph('- plotly: interaktive HTML-Visualisierungen', style='List Bullet')
doc.add_paragraph('- Pillow: Bildverarbeitung für Tkinter', style='List Bullet')
doc.add_paragraph('- python-docx: Erzeugung von Word-Dokumenten', style='List Bullet')

add_heading('9. Anmerkung', level=1)
doc.add_paragraph('Diese Dokumentation wurde automatisch aus der aktuellen Projektarchitektur erstellt. Neue Änderungen in den Modulen sollten hier angepasst werden, damit die Dokumentation aktuell bleibt.')

doc.save(filename)
print(f'Document created: {filename}')
